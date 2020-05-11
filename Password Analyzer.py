# Strong Password detection wth Regexes
# By Devin Glispy
import re
import Plot
import matplotlib.pyplot as plt
import numpy as np
from scipy import stats
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.table import WD_TABLE_ALIGNMENT

with open('Pass10k.txt', encoding="utf-8") as file:
    data = file.read().replace('\n', ',')

data = data.split(',')
print(data)
NumData = 9999  # The default file has 10k passwords feel free to try this out on other password sets
data = data[:NumData]


# START GRADING FUNCTION DONE
def passwordStrength(passwordText, Isfirst):
    GradeVal = 0
    # Strength Checks
    charWeakRegex = re.compile(r'\S{6,}')  # Check if password has atleast 6 characters
    charNormRegex = re.compile(r'\S{9,}')  # Check if password has atleast 9 characters
    charStronRegex = re.compile(r'\S{12,}')  # Check if password has atleast 12 characters
    lowerRegex = re.compile(r'[a-z]+')  # Check if at least one lowercase letter
    upperRegex = re.compile(r'[A-Z]+')  # Check if atleast one upper case letter
    digitRegex = re.compile(r'[0-9]+')  # Check if at least one digit.
    symbolRegex = re.compile(r'[-!$%^&*()_+|~=`{}\[\]:";<>?,./]+')  # Check if at least one symbol.

    # This gives us the Length estimate

    if charWeakRegex.findall(
            passwordText):
        GradeVal = GradeVal + 1
        if Isfirst:
            print('Password contains at least 7 characters')
    else:
        if Isfirst:
            print('Password contains 6 or less characters')
        GradeVal = GradeVal - 1
    if charNormRegex.findall(
            passwordText):
        GradeVal = GradeVal + 2
        if Isfirst:
            print('Password contains at least 9 characters')
    if charStronRegex.findall(
            passwordText):
        GradeVal = GradeVal + 3
        if Isfirst:
            print('Password contains at least 12 characters')

    # Checks for specific Characters
    if lowerRegex.findall(
            passwordText):  # Checks if the password contains a lowercase character and returns a message
        GradeVal = GradeVal + 1
        if Isfirst:
            print('Password contains at least one lowercase letter')

    if upperRegex.findall(
            passwordText):  # Checks if the password contains an uppercase character and returns a message
        GradeVal = GradeVal + 1
        if Isfirst:
            print('Password contains at least one uppercase letter')

    if digitRegex.findall(
            passwordText):  # Checks if the password contains a digit character and returns a message
        GradeVal = GradeVal + 2
        if Isfirst:
            print('Password contains at least one digit character')

    if symbolRegex.findall(
            passwordText):  # Checks if the password contains a digit character and returns a message
        GradeVal = GradeVal + 3
        if Isfirst:
            print('Password contains at least one symbol character')

    return GradeVal  # break out of function.


# GRADING FUNCTION DONE

# PARETO CHART MAKING START

labels = ['1', '2', '3', '4', '5',
          '6', '7', '8', '9', '10', '11', '12', '13']
count = [0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0, 0]

sum1 = 0
Grades = []
cData1 = []
pwd = ''
Grade = []
first = True
for pw in data:

    if first:
        g = passwordStrength(pw, True)
        Grade = g
        pwd = pw
    else:
        g = passwordStrength(pw, False)
    Grades.append(g)
    if g < 13:

        count[g] = count[g] + 1
    else:
        count[12] = count[12] + 1

    first = False
first = True
if first:
    print("pw")
    print(pwd)
    print('Printing Grade')
    print(Grade)
    first = False
first = True
for x in count:
    # Grades.append(str([passwordStrength(pw)]))
    sum1 = sum1 + x
    cData1.append(sum1)
    first = False
pert = stats.percentileofscore(Grades, 13)
print("This grade is in the "+str(pert)+" percentile.")
plot1 = Plot.plot("passwords-2")
plt.title("Hacked passwords by grade")
plt.xlabel("Password Grade")
plt.ylabel("Count")
fig1 = plot1.ax1.bar(labels, count, color='green', edgecolor='black')


def labelBars(bars):
    ###Attach a text label above each bar displaying its height###
    for bar in bars:
        height = bar.get_height()
        plot1.ax1.annotate('{}'.format(height), xy=(bar.get_x() + bar.get_width() / 2, height),
                           xytext=(0, 4), textcoords="offset points", ha='center', va='bottom')


labelBars(fig1)

plot1.ax1.plot(labels, cData1, color="#000000", linestyle="dotted", marker="x")

# PARETO CHART MAKING DONE

# DOCUMENT MAKING START
plot1.save(tight1="tight")
imageFs1 = "passwords-2-01.png"
tableList1 = [["1"], ["2"], ["3"], ["4"], ["5"], ["6"], ["7"], ["8"], ["9"], ["10"], ["11"], ["12"], ["13"]]
x = 0
for num in tableList1:
    num.append(count[x])
    x = x + 1

doc1 = Document()
doc1.add_heading("Final Project: Password Analysis", 0)
listlen1 = len(tableList1)
clen1 = len(tableList1[0])
slen1 = 4  # NUM SECTIONS IN DOC
docFs1 = "./Write-up.docx"
for pos1 in range(1, slen1):

    if pos1 == 1:
        heading1 = "Section {0:d}: Introductions".format(pos1)
        doc1.add_heading(heading1, 1)
        para1 = doc1.add_paragraph(
            "This is an analysis of how well my grading parameters can assess the top 10,000 breached passwords." +
            " I have chosen to use the list provided by www.haveibeenpwed.com for reference.  I think this will give"
            " us a enough data to make accurate assumptions about weak passwords in general." +
            " The first password in the list is a password for you to change and play with." +
            " Currently it is '" + str(data[0]) + "' and on our point-based grading scale that goes from 1 to 13."
            " Its grade was " + str(Grade) + ".")

    elif pos1 == 2:
        heading1 = "Section {0:d}: Table Data".format(pos1)
        doc1.add_heading(heading1, 1)
        

        para1 = doc1.add_paragraph(
            "Here is a table that shows the grades gotten by the top 10,000 passwords." +
            " Now, Because your score was " + str(Grade) + " you actually scored in the " +
            "{0:.3f} percentile.".format(pert))

        print("Insert table at section {0:d}.".format(pos1))
        table1 = doc1.add_table(rows=listlen1, cols=clen1, style="Table Grid")
        for rpos1 in range(0, listlen1):
            for cpos1 in range(0, clen1):
                cell1 = table1.cell(rpos1, cpos1)
                text1 = tableList1[rpos1][cpos1]
                cell1.text = str(text1)
                cell1.paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
                # print("\t( {0:d} , {1:d} ) \"{2:s}\"".format(rpos1, cpos1, str(text1)))

    elif pos1 == 3:
        doc1.add_page_break()
        heading1 = "Section {0:d}: Chart Data".format(pos1)
        doc1.add_heading(heading1, 1)
        
        PointsChart = [
            ('less than 6 charaters', 0),
            ("more than 6 charaters", 1),
            ("more than 9 charaters", 3),
            ("more than 12 charaters", 5),
            ("uses a lowercase letter", 1),
            ("uses a uppercase letter", 1),
            ("uses a number", 2),
            ("uses a Symbol ", 3)
        ]
        points = [0,1,3,5,1,1,2,3]
        table = doc1.add_table(rows=1, cols=2, style="Table Grid")

        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Description'
        hdr_cells[1].text = 'Points'

        for desc, points in PointsChart:
            row_cells = table.add_row().cells
            row_cells[0].text = desc
            row_cells[1].text = str(points)




        print("Insert image at section {0:d}.".format(pos1))
        doc1.add_picture(imageFs1, width=Inches(6.0))
        

doc1.save(docFs1)

# DOCUMENT MAKING DONE
input("Press Enter to continue...")